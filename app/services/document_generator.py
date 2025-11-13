from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
import os

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def generate_word_document(english_text, chinese_translation, vocabulary_list, output_path):
    """
    生成包含英文原文、中文翻译和词汇表的Word文档
    
    参数:
    english_text (str): 原始英文文本
    chinese_translation (str): 中文翻译文本
    vocabulary_list (list): 词汇表列表，每个元素是包含english、chinese和explanation的字典
    output_path (str): 输出文档的保存路径
    
    返回:
    bool: 文档生成是否成功
    """
    try:
        # 创建一个新的Word文档
        doc = Document()
        
        # 设置文档样式
        setup_document_styles(doc)
        
        # 添加标题
        title = doc.add_heading('翻译结果文档', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加英文原文部分
        doc.add_heading('一、英文原文', level=1)
        english_section = doc.add_paragraph(english_text)
        
        # 添加中文翻译部分
        doc.add_heading('二、中文翻译', level=1)
        chinese_section = doc.add_paragraph(chinese_translation)
        
        # 添加词汇表部分（如果有）
        if vocabulary_list:
            doc.add_heading('三、专业词汇表', level=1)
            
            # 创建词汇表表格
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            # 设置表头
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '英文术语'
            hdr_cells[1].text = '中文翻译'
            hdr_cells[2].text = '术语解释'
            
            # 设置表头样式
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加词汇表内容
            for vocab in vocabulary_list:
                row_cells = table.add_row().cells
                row_cells[0].text = vocab.get('english', '')
                row_cells[1].text = vocab.get('chinese', '')
                row_cells[2].text = vocab.get('explanation', '')
            
            # 添加一个空段落作为分隔
            doc.add_paragraph()
        
        # 添加页脚
        footer_text = "此文档由 xxx 翻译助手自动生成"
        for section in doc.sections:
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = footer_text
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 保存文档
        # 确保输出目录存在
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        doc.save(output_path)
        logger.info(f"Word文档已成功生成并保存至: {output_path}")
        return True
        
    except Exception as e:
        logger.error(f"Word文档生成失败: {str(e)}")
        raise Exception(f"文档生成失败: {str(e)}")

def setup_document_styles(doc):
    """
    设置文档样式
    
    参数:
    doc: Document对象
    """
    # 获取样式集合
    styles = doc.styles
    
    # 设置正文样式
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Microsoft YaHei'
    normal_font.size = Pt(12)
    
    # 设置标题样式
    heading_styles = ['Heading 1', 'Heading 2', 'Heading 3']
    heading_font_sizes = [Pt(16), Pt(14), Pt(12)]
    
    for i, style_name in enumerate(heading_styles):
        if style_name in styles:
            heading_style = styles[style_name]
            heading_font = heading_style.font
            heading_font.name = 'Microsoft YaHei'
            heading_font.size = heading_font_sizes[i]
            heading_font.bold = True