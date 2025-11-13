import unittest
import os
import sys
import shutil
from docx import Document

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from app.services.document_generator import generate_word_document

class TestDocumentGenerator(unittest.TestCase):
    
    def setUp(self):
        """测试前设置"""
        self.test_dir = "test_downloads"
        os.makedirs(self.test_dir, exist_ok=True)
        self.test_file = os.path.join(self.test_dir, "test_doc.docx")
    
    def tearDown(self):
        """测试后清理"""
        if os.path.exists(self.test_dir):
            shutil.rmtree(self.test_dir)
    
    def test_generate_word_document_basic(self):
        """测试基本文档生成功能"""
        # 准备测试数据
        original_text = "This is a test document for translation."
        translated_text = "这是一个用于翻译的测试文档。"
        vocabulary = []
        
        # 生成文档
        generate_word_document(original_text, translated_text, vocabulary, self.test_file)
        
        # 验证文档是否创建
        self.assertTrue(os.path.exists(self.test_file))
        
        # 验证文档内容
        doc = Document(self.test_file)
        paragraphs = [p.text for p in doc.paragraphs]
        
        # 检查是否包含必要的段落
        self.assertIn("一、英文原文", paragraphs)
        self.assertIn(original_text, paragraphs)
        self.assertIn("二、中文翻译", paragraphs)
        self.assertIn(translated_text, paragraphs)
    
    def test_generate_word_document_with_vocabulary(self):
        """测试带词汇表的文档生成"""
        # 准备测试数据
        original_text = "Machine Learning is important."
        translated_text = "机器学习很重要。"
        vocabulary = [{
            "english": "Machine Learning",
            "chinese": "机器学习",
            "explanation": "人工智能的一个分支"
        }]
        
        # 生成文档
        generate_word_document(original_text, translated_text, vocabulary, self.test_file)
        
        # 验证文档是否创建
        self.assertTrue(os.path.exists(self.test_file))
        
        # 验证文档内容
        doc = Document(self.test_file)
        paragraphs = [p.text for p in doc.paragraphs]
        
        # 检查是否包含词汇表标题
        self.assertIn("三、专业词汇表", paragraphs)
        
        # 检查表格内容
        tables = doc.tables
        self.assertEqual(len(tables), 1)
        table = tables[0]
        self.assertEqual(table.cell(0, 0).text, "英文术语")
        self.assertEqual(table.cell(0, 1).text, "中文翻译")
        self.assertEqual(table.cell(0, 2).text, "术语解释")
        self.assertEqual(table.cell(1, 0).text, "Machine Learning")

if __name__ == '__main__':
    unittest.main()